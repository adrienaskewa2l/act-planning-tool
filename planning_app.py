#!/usr/bin/env python3
"""
ACT Conference 2026 – Planning Interactif
==========================================
Lancez:  python3 planning_app.py
Puis ouvrez: http://localhost:5001
"""

import os, json, io
from flask import Flask, jsonify, request, send_file, Response
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources={
    r"/api/*": {
        "origins": [
            "http://localhost:5001",
            "http://localhost:5002",
            "http://127.0.0.1:5001",
            "http://127.0.0.1:5002"
        ]
    }
})
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, "schedule_data.json")
JS_FILE = os.path.join(BASE_DIR, "app.js")

# ─────────────────────────────────────────────────────────────────
# DONNÉES INITIALES (depuis le PDF + Running Sheet)
# ─────────────────────────────────────────────────────────────────
DEFAULT_SCHEDULE = {
    "version": "1.0",
    "event": "ACT Conference 2026",
    "lieu": "Espace 140, Rillieux la Pape",
    "dresscode": "Bracelet bleu pour la Dream Team",
    "installation_dates": "Jeudi 14 mai 2026",
    "conference_dates": "Vendredi 15 – Dimanche 17 mai 2026",
    "attendance": {
        "vendredi": "",
        "samedi": "",
        "dimanche": ""
    },
    "types": {
        "LOUANGE": {"label": "Louange", "color": "#C39BD3"},
        "PREDICATION": {"label": "Prédication", "color": "#82E0AA"},
        "REPAS": {"label": "Repas", "color": "#FAD7A0"},
        "LOGISTIQUE": {"label": "Logistique", "color": "#AED6F1"},
        "REPETITION": {"label": "Répétition", "color": "#D7BDE2"},
        "REUNION": {"label": "Réunion / Brief", "color": "#D5F5E3"},
        "BOOST": {"label": "Boost", "color": "#F1948A"},
        "STE_CENE": {"label": "Ste Cène", "color": "#BB8FCE"},
        "PRIERE": {"label": "Prière", "color": "#7FB3D3"},
        "INTERCESSION": {"label": "Intercession", "color": "#D6EAF8"},
        "FUNDRAISING": {"label": "Levée de fonds", "color": "#F9E79F"},
        "TED_PITCH": {"label": "Pitch / TED", "color": "#A9DFBF"},
        "ANNONCES": {"label": "Annonces", "color": "#A9DFBF"},
        "TECHNIQUE": {"label": "Technique", "color": "#D5D8DC"}
    },
    "days": [
        {
            "id": "jeudi",
            "name": "JEUDI 14 MAI",
            "date": "2026-05-14",
            "sessions": [
                {
                    "id": "jeudi_install",
                    "name": "Installation",
                    "events": [
                        {"id": "j001", "time": "07:00", "duration": 60,
                         "title": "Arrivée St Joseph – Chargement matériel",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Team Logistique"],
                         "details": "Arrivée à St Joseph pour chargement du matériel"},
                        {"id": "j002", "time": "08:00", "duration": 15,
                         "title": "Départ vers E140",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Team Logistique"],
                         "details": "Départ à E140. Adrien part en avance."},
                        {"id": "j003", "time": "08:15", "duration": 15,
                         "title": "Installation café/buffet",
                         "type": "LOGISTIQUE", "color": "#FAD7A0",
                         "teams": ["Team Hospitalité"],
                         "details": "Installation café/buffet (en route vers E140)"},
                        {"id": "j004", "time": "08:30", "duration": 30,
                         "title": "Arrivée E140 – Café/Buffet prêt",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Team Hospitalité", "Team Logistique"],
                         "details": "Arrivée à E140. Café/Buffet prêt sur place. Arrivée des responsables."},
                        {"id": "j005", "time": "09:00", "duration": 210,
                         "title": "Brief organisation + Installations multiples",
                         "type": "REUNION", "color": "#D5F5E3",
                         "teams": ["Team Coordo", "Team Logistique", "Team Kids", "Team Bébés",
                                   "Team Snacking", "Team Communication", "Team Médias",
                                   "Team Bible School", "Team Festival", "Team Librairie"],
                         "details": "Briefing organisation + lancement. Installation des équipes (Logistique, Kids, Bébés, Ménage). Préparation snacking (ados + jeunes). Communication RS stories. Installation des stands (Bible School, Festival, Librairie)."},
                        {"id": "j006", "time": "12:30", "duration": 60,
                         "title": "Repas tirés des sacs",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": [], "details": "Repas tirés des sacs"},
                        {"id": "j007", "time": "13:30", "duration": 30,
                         "title": "Derniers ajustements installation",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Toutes les teams"],
                         "details": "Derniers ajustements installation"},
                        {"id": "j008", "time": "14:00", "duration": 90,
                         "title": "Arrivée louange + technique + installation",
                         "type": "REPETITION", "color": "#D7BDE2",
                         "teams": ["Team Louange", "Team Son"],
                         "details": "Arrivée et installation des musiciens (disposition...)"},
                        {"id": "j009", "time": "15:30", "duration": 150,
                         "title": "Balances et répétitions",
                         "type": "REPETITION", "color": "#D7BDE2",
                         "teams": ["Team Louange", "Team Son"],
                         "details": "Balances et répétitions"},
                        {"id": "j010", "time": "18:00", "duration": 15,
                         "title": "Fin de journée – Départ",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Toutes les teams"],
                         "details": "Fin de journée d'installation. Départ de tout le monde."}
                    ]
                }
            ]
        },
        {
            "id": "vendredi",
            "name": "VENDREDI 15 MAI",
            "date": "2026-05-15",
            "sessions": [
                {
                    "id": "ven_prep",
                    "name": "Préparation",
                    "events": [
                        {"id": "v001", "time": "08:30", "duration": 30,
                         "title": "Arrivée à E140",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Toutes les teams", "Toute la Dream Team ACT"],
                         "details": "Arrivée à E140 – Toutes les teams, toute la Dream Team ACT"},
                        {"id": "v002", "time": "09:00", "duration": 30,
                         "title": "Briefing organisation + lancement",
                         "type": "REUNION", "color": "#D5F5E3",
                         "teams": ["Team Coordo"],
                         "details": "Briefing organisation + lancement"},
                        {"id": "v003", "time": "09:30", "duration": 150,
                         "title": "Répétitions + Préparations multiples",
                         "type": "REPETITION", "color": "#D7BDE2",
                         "teams": ["Team Louange", "Team Son", "Team Snacking",
                                   "Team Médias Capture", "Team Médias Projection",
                                   "Team Sécurité", "Team Parking", "Team Kids",
                                   "Team Bébés", "Team Hospitalité", "Fil Rouge"],
                         "details": "Répétitions de louange. Dernières préparations du snacking. Ajustements techniques, sécurité."},
                        {"id": "v004", "time": "12:00", "duration": 90,
                         "title": "Ménage + Repas",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": ["Team E140", "Team Snacking"],
                         "details": "Ménage – Team E140. Repas tirés des sacs OU snacking déjà opérationnel – Team Snacking ?"},
                        {"id": "v005", "time": "13:30", "duration": 30,
                         "title": "Méga-Boost – Toute la Dream Team",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost"],
                         "details": "Méga-Boost pour toute la Dream Team ACT – Team Boost, mot d'Adrien, mot du Fil Rouge, mot de Matt ?"},
                        {"id": "v006", "time": "14:00", "duration": 30,
                         "title": "Répétitions de louange",
                         "type": "REPETITION", "color": "#D7BDE2",
                         "teams": ["Team Louange", "Team Son"],
                         "details": "Répétitions de louange"},
                        {"id": "v007", "time": "14:30", "duration": 75,
                         "title": "Accueil opérationnel",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Team Accueil", "Team Snacking", "Team Festival",
                                   "Team Bible School", "Team Librairie", "Team Hospitalité"],
                         "details": "Accueil opérationnel : entrée dans le hall, stands ouverts (Snacking, Librairie, Festival, Bible School). Portes de la salle principale fermées. Scan des billets. Orateurs principaux → Green Room."}
                    ]
                },
                {
                    "id": "ven_ouverture",
                    "name": "Célébration d'ouverture",
                    "events": [
                        {"id": "v010", "time": "15:45", "duration": 10,
                         "title": "Boost",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost"],
                         "details": "Boost"},
                        {"id": "v011", "time": "15:55", "duration": 5,
                         "title": "Point Fil Rouge",
                         "type": "REUNION", "color": "#D5F5E3",
                         "teams": ["Fil Rouge", "Équipes de service"],
                         "details": "Point Fil Rouge – Fil Rouge + Équipes de service Célébration d'ouverture"},
                        {"id": "v012", "time": "16:00", "duration": 25,
                         "title": "Intercession + Ouverture des portes",
                         "type": "INTERCESSION", "color": "#D6EAF8",
                         "teams": ["Team Prière", "Team Accueil"],
                         "details": "Intercession – Josias, Team Prière, équipiers volontaires (hors service). Ouverture des portes de la salle principale – Team Accueil."},
                        {"id": "v013", "time": "16:25", "duration": 5,
                         "title": "Lancement vidéo pre-roll + countdown",
                         "type": "TECHNIQUE", "color": "#D5D8DC",
                         "teams": ["Team Médias Projection"],
                         "details": "Lancement vidéo pre-roll + countdown – Team Médias Projection"},
                        {"id": "v014", "time": "16:30", "duration": 5,
                         "title": "1er chant",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange", "Team Son", "Team Médias Projection",
                                   "Team Médias Capture", "Team Communication"],
                         "details": "Premier chant. Photos et vidéos en continu jusqu'à la fin de la célébration."},
                        {"id": "v015", "time": "16:35", "duration": 25,
                         "title": "Mot d'accueil – MC",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": ["MC"],
                         "details": "Mot d'accueil – MC"},
                        {"id": "v016", "time": "17:00", "duration": 30,
                         "title": "Introduction Johannes + Annonces WE",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": ["Johannes"],
                         "details": "Introduction Johannes, Annonces importantes pour le WE"},
                        {"id": "v017", "time": "17:30", "duration": 30,
                         "title": "Louange",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange"},
                        {"id": "v018", "time": "18:00", "duration": 75,
                         "title": "Repas",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": ["Team Snacking"],
                         "details": "Repas"},
                        {"id": "v019", "time": "19:15", "duration": 15,
                         "title": "Boost + point fil rouge",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost", "Fil Rouge"],
                         "details": "Boost + point fil rouge"},
                        {"id": "v020", "time": "19:30", "duration": 30,
                         "title": "Préparation soirée",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": [], "details": ""},
                        {"id": "v021", "time": "20:00", "duration": 45,
                         "title": "Louange",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange"},
                        {"id": "v022", "time": "20:45", "duration": 45,
                         "title": "Prédication Dennis",
                         "type": "PREDICATION", "color": "#82E0AA",
                         "teams": ["Dennis"],
                         "details": "Prédication Dennis"},
                        {"id": "v023", "time": "21:30", "duration": 30,
                         "title": "Louange + ministère",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange, ministère…"},
                        {"id": "v024", "time": "22:00", "duration": 30,
                         "title": "Levée de fonds – 10K pour ACT et BS",
                         "type": "FUNDRAISING", "color": "#F9E79F",
                         "teams": [], "details": "Levée de fonds : 10K pour ACT et BS"}
                    ]
                }
            ]
        },
        {
            "id": "samedi",
            "name": "SAMEDI 16 MAI",
            "date": "2026-05-16",
            "sessions": [
                {
                    "id": "sam_matin",
                    "name": "Matinée",
                    "events": [
                        {"id": "s001", "time": "08:30", "duration": 15,
                         "title": "Arrivée Dream Team",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Toute la Dream Team"],
                         "details": "Arrivée Dream Team"},
                        {"id": "s002", "time": "08:45", "duration": 45,
                         "title": "Boost + point fil rouge",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost", "Fil Rouge"],
                         "details": "Boost + point fil rouge"},
                        {"id": "s003", "time": "09:30", "duration": 30,
                         "title": "Louange",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange"},
                        {"id": "s004", "time": "10:00", "duration": 30,
                         "title": "Prédication Dennis",
                         "type": "PREDICATION", "color": "#82E0AA",
                         "teams": ["Dennis"],
                         "details": "Prédication Dennis"},
                        {"id": "s005", "time": "10:30", "duration": 15,
                         "title": "Ste Cène",
                         "type": "STE_CENE", "color": "#BB8FCE",
                         "teams": ["Johannes"],
                         "details": "Ste Cène conduite par Johannes"},
                        {"id": "s006", "time": "10:45", "duration": 15,
                         "title": "Temps de prière",
                         "type": "PRIERE", "color": "#7FB3D3",
                         "teams": ["Johannes"],
                         "details": "Temps de prière (Johannes)"},
                        {"id": "s007", "time": "11:00", "duration": 60,
                         "title": "Louange + ministère",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange, ministère…"},
                        {"id": "s008", "time": "12:00", "duration": 150,
                         "title": "Repas",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": [], "details": "Repas"}
                    ]
                },
                {
                    "id": "sam_apresmidi",
                    "name": "Après-midi",
                    "events": [
                        {"id": "s010", "time": "14:30", "duration": 15,
                         "title": "Pitchs courts format TED",
                         "type": "TED_PITCH", "color": "#A9DFBF",
                         "teams": [], "details": "Pitchs courts format TED"},
                        {"id": "s011", "time": "14:45", "duration": 15,
                         "title": "Nouvelles églises + implantations",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": [], "details": "Nouvelles églises, implantations…"},
                        {"id": "s012", "time": "15:00", "duration": 15,
                         "title": "Podcast Johannes",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": ["Johannes"],
                         "details": "Podcast Johannes"},
                        {"id": "s013", "time": "15:15", "duration": 15,
                         "title": "Deese – intervention type TED",
                         "type": "TED_PITCH", "color": "#A9DFBF",
                         "teams": ["Deese"],
                         "details": "Deese : intervention type TED"},
                        {"id": "s014", "time": "15:30", "duration": 15,
                         "title": "Transition",
                         "type": "LOGISTIQUE", "color": "#D5D8DC",
                         "teams": [], "details": ""},
                        {"id": "s015", "time": "15:45", "duration": 15,
                         "title": "Boost + point fil rouge",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost", "Fil Rouge"],
                         "details": "Boost + point fil rouge"},
                        {"id": "s016", "time": "16:00", "duration": 15,
                         "title": "Écoles bibliques + festivals + témoignages",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": [], "details": "Écoles bibliques, festivals, témoignages"},
                        {"id": "s017", "time": "16:15", "duration": 15,
                         "title": "Vidéos news églises SOS",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": ["Team Médias Projection"],
                         "details": "Vidéos news églises SOS"},
                        {"id": "s018", "time": "16:30", "duration": 30,
                         "title": "Présentation network prière 24/24",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": [], "details": "Présentation network prière 24/24"},
                        {"id": "s019", "time": "17:00", "duration": 30,
                         "title": "À compléter",
                         "type": "ANNONCES", "color": "#D5D8DC",
                         "teams": [], "details": ""},
                        {"id": "s020", "time": "17:30", "duration": 60,
                         "title": "Intervention 3",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": [], "details": "Intervention 3"},
                        {"id": "s021", "time": "18:30", "duration": 45,
                         "title": "Repas",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": [], "details": "Repas"}
                    ]
                },
                {
                    "id": "sam_soiree",
                    "name": "Soirée",
                    "events": [
                        {"id": "s030", "time": "19:15", "duration": 15,
                         "title": "Boost + point fil rouge",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost", "Fil Rouge"],
                         "details": "Boost + point fil rouge"},
                        {"id": "s031", "time": "19:30", "duration": 30,
                         "title": "Préparation soirée",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": [], "details": ""},
                        {"id": "s032", "time": "20:00", "duration": 30,
                         "title": "Louange Résistance",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange Résistance"},
                        {"id": "s033", "time": "20:30", "duration": 60,
                         "title": "Prédication Johannes",
                         "type": "PREDICATION", "color": "#82E0AA",
                         "teams": ["Johannes"],
                         "details": "Prédication Johannes"},
                        {"id": "s034", "time": "21:30", "duration": 30,
                         "title": "Temps de prière",
                         "type": "PRIERE", "color": "#7FB3D3",
                         "teams": ["Johannes"],
                         "details": "Temps de prière conduit par Johannes"},
                        {"id": "s035", "time": "22:00", "duration": 15,
                         "title": "Ministère",
                         "type": "PRIERE", "color": "#7FB3D3",
                         "teams": [], "details": "Ministère…"},
                        {"id": "s036", "time": "22:15", "duration": 15,
                         "title": "Levée de fonds – 10K pour ACT et BS",
                         "type": "FUNDRAISING", "color": "#F9E79F",
                         "teams": [], "details": "Levée de fonds : 10K pour ACT et BS"}
                    ]
                }
            ]
        },
        {
            "id": "dimanche",
            "name": "DIMANCHE 17 MAI",
            "date": "2026-05-17",
            "sessions": [
                {
                    "id": "dim_matin",
                    "name": "Matinée – Clôture",
                    "events": [
                        {"id": "d001", "time": "08:30", "duration": 15,
                         "title": "Arrivée Dream Team",
                         "type": "LOGISTIQUE", "color": "#AED6F1",
                         "teams": ["Toute la Dream Team"],
                         "details": "Arrivée Dream Team"},
                        {"id": "d002", "time": "08:45", "duration": 15,
                         "title": "Boost + point fil rouge (attention Ste Cène)",
                         "type": "BOOST", "color": "#F1948A",
                         "teams": ["Team Boost", "Fil Rouge"],
                         "details": "Boost + point fil rouge (attention Ste Cène)"},
                        {"id": "d003", "time": "09:00", "duration": 30,
                         "title": "Intercession + Ouverture des portes",
                         "type": "INTERCESSION", "color": "#D6EAF8",
                         "teams": ["Team Prière", "Team Accueil"],
                         "details": "Intercession / Ouverture des portes"},
                        {"id": "d004", "time": "09:30", "duration": 60,
                         "title": "Louange",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange"},
                        {"id": "d005", "time": "10:30", "duration": 15,
                         "title": "Ste Cène",
                         "type": "STE_CENE", "color": "#BB8FCE",
                         "teams": ["Johannes"],
                         "details": "Ste Cène conduite par Johannes"},
                        {"id": "d006", "time": "10:45", "duration": 15,
                         "title": "Temps de prière",
                         "type": "PRIERE", "color": "#7FB3D3",
                         "teams": ["Johannes"],
                         "details": "Temps de prière (Johannes)"},
                        {"id": "d007", "time": "11:00", "duration": 60,
                         "title": "Prédication Dennis",
                         "type": "PREDICATION", "color": "#82E0AA",
                         "teams": ["Dennis"],
                         "details": "Prédication Dennis (option Mamadou si présent)"},
                        {"id": "d008", "time": "12:00", "duration": 15,
                         "title": "Louange + ministère",
                         "type": "LOUANGE", "color": "#C39BD3",
                         "teams": ["Team Louange"],
                         "details": "Louange, ministère"},
                        {"id": "d009", "time": "12:15", "duration": 15,
                         "title": "Levée de fonds – 10K pour ACT et BS",
                         "type": "FUNDRAISING", "color": "#F9E79F",
                         "teams": [], "details": "Levée de fonds : 10K pour ACT et BS"},
                        {"id": "d010", "time": "12:30", "duration": 15,
                         "title": "Clôture conférence + annonces",
                         "type": "ANNONCES", "color": "#A9DFBF",
                         "teams": [], "details": "Clôture conférence, annonces"},
                        {"id": "d011", "time": "12:45", "duration": 75,
                         "title": "Repas",
                         "type": "REPAS", "color": "#FAD7A0",
                         "teams": [], "details": "Repas"}
                    ]
                }
            ]
        }
    ]
}


def default_copy():
    return json.loads(json.dumps(DEFAULT_SCHEDULE))


def normalize_schedule(data):
    schedule = data or default_copy()
    schedule.setdefault("version", DEFAULT_SCHEDULE["version"])
    schedule.setdefault("event", DEFAULT_SCHEDULE["event"])
    schedule.setdefault("lieu", DEFAULT_SCHEDULE["lieu"])
    schedule.setdefault("dresscode", DEFAULT_SCHEDULE["dresscode"])
    schedule.setdefault("installation_dates", DEFAULT_SCHEDULE["installation_dates"])
    schedule.setdefault("conference_dates", DEFAULT_SCHEDULE["conference_dates"])
    schedule.setdefault("attendance", {})
    schedule.setdefault("types", {})

    for key, value in DEFAULT_SCHEDULE["attendance"].items():
        schedule["attendance"].setdefault(key, value)

    for key, value in DEFAULT_SCHEDULE["types"].items():
        schedule["types"].setdefault(key, value.copy())

    for day in schedule.get("days", []):
        day.setdefault("id", "")
        day.setdefault("name", "")
        day.setdefault("date", "")
        day.setdefault("sessions", [])
        for session in day.get("sessions", []):
            session.setdefault("id", "")
            session.setdefault("name", "Session")
            session.setdefault("events", [])
            session["events"].sort(key=lambda ev: ev.get("time", "00:00"))
            for event in session["events"]:
                event.setdefault("id", "")
                event.setdefault("time", "09:00")
                event.setdefault("duration", 30)
                event.setdefault("title", "")
                event.setdefault("type", "ANNONCES")
                event.setdefault("color", "#A9DFBF")
                event.setdefault("teams", [])
                event.setdefault("details", "")
                if event["type"] not in schedule["types"]:
                    schedule["types"][event["type"]] = {
                        "label": event["type"].replace("_", " ").title(),
                        "color": event.get("color") or "#A9DFBF"
                    }
    return schedule

# ─────────────────────────────────────────────────────────────────
# HTML TEMPLATE
# ─────────────────────────────────────────────────────────────────
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ACT 2026 – Planning</title>
<style>
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  :root {
    --header-h: 64px;
    --axis-w: 54px;
    --col-w: 230px;
    --px-per-min: 1.6;
    --start-h: 7;
    --end-h: 23;
    --green: #0F4C3A;
    --pink: #CF5AFD;
    --font: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
  }
  body { font-family: var(--font); background: #f0f2f5; color: #1a1a1a; }

  /* ── TOOLBAR ── */
  .toolbar {
    position: fixed; top: 0; left: 0; right: 0; height: var(--header-h);
    background: var(--green); color: white;
    display: flex; align-items: center; gap: 12px; padding: 0 20px;
    z-index: 100; box-shadow: 0 2px 8px rgba(0,0,0,.3);
  }
  .toolbar h1 { font-size: 16px; font-weight: 600; flex: 1; }
  .toolbar .logo { font-size: 22px; }
  .btn {
    padding: 7px 14px; border: none; border-radius: 8px; cursor: pointer;
    font-size: 13px; font-weight: 600; transition: opacity .15s;
  }
  .btn:hover { opacity: .85; }
  .btn-save { background: #fff; color: var(--green); }
  .btn-docx { background: #4CAF50; color: white; }
  .btn-sync { background: #F9E79F; color: #174032; }
  .btn-add  { background: rgba(255,255,255,.18); color: white; border: 1px solid rgba(255,255,255,.4); }
  .toast {
    position: fixed; bottom: 20px; right: 20px; background: #333; color: white;
    padding: 10px 18px; border-radius: 8px; font-size: 13px;
    opacity: 0; transition: opacity .3s; z-index: 200; pointer-events: none;
  }
  .toast.show { opacity: 1; }

  /* ── LEGEND ── */
  .legend {
    position: fixed; left: 0; top: var(--header-h); bottom: 0; width: 160px;
    background: white; border-right: 1px solid #e0e0e0;
    overflow-y: auto; padding: 12px 10px;
  }
  .legend h3 { font-size: 11px; text-transform: uppercase; color: #888; margin-bottom: 8px; letter-spacing: .5px; }
  .legend-item {
    display: flex; align-items: center; gap: 6px;
    font-size: 11px; padding: 3px 0; cursor: pointer;
  }
  .legend-label { flex: 1; min-width: 0; }
  .legend-edit {
    border: 0; background: transparent; color: #777; cursor: pointer;
    font-size: 12px; line-height: 1; padding: 3px; border-radius: 4px;
  }
  .legend-edit:hover { background: #f0f0f0; color: var(--green); }
  .legend-add {
    width: 100%; margin: 8px 0 12px; padding: 7px 8px;
    border: 1px dashed #bbb; border-radius: 7px; background: #fafafa;
    color: var(--green); cursor: pointer; font-size: 11px; font-weight: 700;
  }
  .legend-add:hover { background: #f3f7f5; border-color: var(--green); }
  .legend-item.hidden-type { opacity: .35; text-decoration: line-through; }
  .legend-dot { width: 12px; height: 12px; border-radius: 3px; flex-shrink: 0; }

  /* ── MAIN AREA ── */
  .main {
    margin-top: var(--header-h);
    margin-left: 160px;
    overflow-x: auto;
    padding-bottom: 40px;
  }

  /* ── GRID ── */
  .grid-wrapper {
    display: flex;
    min-width: calc(var(--axis-w) + var(--col-w) * 4);
    position: relative;
  }

  /* ── TIME AXIS ── */
  .time-axis {
    width: var(--axis-w);
    flex-shrink: 0;
    position: relative;
    margin-top: 42px;
  }
  .hour-label {
    position: absolute; right: 8px;
    font-size: 11px; color: #888; line-height: 1;
    transform: translateY(-50%);
  }

  /* ── DAY COLUMNS ── */
  .days-row {
    display: flex;
    flex: 1;
  }
  .day-col {
    width: var(--col-w);
    flex-shrink: 0;
    border-left: 1px solid #d0d0d0;
    position: relative;
  }
  .day-header {
    height: 42px;
    background: var(--green);
    color: white;
    display: flex; align-items: center; justify-content: space-between;
    padding: 0 8px;
    font-size: 12px; font-weight: 700;
    position: sticky; top: var(--header-h); z-index: 10;
  }
  .day-header .add-btn {
    background: rgba(255,255,255,.2); border: none; color: white;
    width: 22px; height: 22px; border-radius: 50%;
    cursor: pointer; font-size: 16px; line-height: 20px; text-align: center;
    flex-shrink: 0;
  }
  .day-header .add-btn:hover { background: rgba(255,255,255,.35); }

  .events-area {
    position: relative;
    background: repeating-linear-gradient(
      to bottom,
      transparent,
      transparent calc(var(--px-per-min) * 60px - 1px),
      #e8e8e8 calc(var(--px-per-min) * 60px - 1px),
      #e8e8e8 calc(var(--px-per-min) * 60px)
    );
  }
  .half-line {
    position: absolute; left: 0; right: 0; height: 1px;
    background: #f0f0f0; pointer-events: none;
  }

  /* ── SESSION SEPARATOR ── */
  .session-sep {
    position: absolute; left: 0; right: 0;
    display: flex; align-items: center; gap: 4px;
    pointer-events: none; z-index: 2;
  }
  .session-sep-line { flex: 1; height: 1px; background: var(--pink); opacity: .4; }
  .session-sep-label {
    font-size: 9px; font-weight: 700; color: var(--pink);
    white-space: nowrap; padding: 0 4px; opacity: .8;
  }

  /* ── EVENT BLOCKS ── */
  .event-block {
    position: absolute;
    left: 4px; right: 4px;
    border-radius: 6px;
    padding: 3px 6px 8px 6px;
    cursor: pointer;
    overflow: hidden;
    border-left: 4px solid rgba(0,0,0,.15);
    transition: box-shadow .15s, transform .1s;
    z-index: 3;
  }
  .event-block:hover {
    box-shadow: 0 2px 12px rgba(0,0,0,.2);
    transform: scaleX(1.01);
    z-index: 5;
  }
  .event-block.hidden-type { display: none; }
  .ev-time { font-size: 10px; opacity: .7; font-weight: 600; line-height: 1.2; }
  .ev-title { font-size: 11px; font-weight: 700; line-height: 1.3; margin-top: 1px; }
  .ev-teams { font-size: 9px; color: var(--pink); margin-top: 2px; line-height: 1.3; }
  .resize-handle {
    position: absolute; bottom: 0; left: 0; right: 0;
    height: 6px; cursor: ns-resize;
    background: rgba(0,0,0,.08);
    border-radius: 0 0 4px 4px;
  }
  .resize-handle:hover { background: rgba(0,0,0,.2); }

  /* ── MODAL ── */
  .modal-overlay {
    position: fixed; inset: 0; background: rgba(0,0,0,.4);
    display: flex; align-items: center; justify-content: center;
    z-index: 150; padding: 20px;
  }
  .modal-overlay.hidden { display: none; }
  .modal-card {
    background: white; border-radius: 12px;
    width: 520px; max-width: 100%; max-height: 90vh;
    overflow-y: auto;
    box-shadow: 0 8px 40px rgba(0,0,0,.2);
    padding: 24px;
  }
  .modal-card h2 { font-size: 17px; color: var(--green); margin-bottom: 18px; }
  .form-row { margin-bottom: 14px; }
  .form-row label {
    display: block; font-size: 11px; font-weight: 600;
    color: #555; margin-bottom: 4px; text-transform: uppercase; letter-spacing: .5px;
  }
  .form-row input, .form-row select, .form-row textarea {
    width: 100%; padding: 8px 10px; border: 1.5px solid #d0d0d0;
    border-radius: 7px; font-size: 13px; font-family: var(--font);
    outline: none; transition: border-color .2s;
  }
  .form-row input:focus, .form-row select:focus, .form-row textarea:focus {
    border-color: var(--green);
  }
  .form-row textarea { min-height: 70px; resize: vertical; }
  .form-row-2 { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-bottom: 14px; }
  .color-row { display: flex; align-items: center; gap: 8px; }
  .color-row input[type=color] {
    width: 40px; height: 36px; padding: 2px; border-radius: 6px; cursor: pointer;
  }
  .type-color-presets { display: flex; gap: 6px; flex-wrap: wrap; margin-top: 6px; }
  .color-preset {
    width: 22px; height: 22px; border-radius: 4px;
    cursor: pointer; border: 2px solid transparent;
    transition: border-color .15s;
  }
  .color-preset:hover, .color-preset.active { border-color: #333; }
  .modal-actions { display: flex; gap: 8px; margin-top: 20px; justify-content: flex-end; }
  .btn-modal { padding: 8px 16px; border: none; border-radius: 7px; cursor: pointer; font-size: 13px; font-weight: 600; }
  .btn-primary { background: var(--green); color: white; }
  .btn-danger  { background: #e53935; color: white; }
  .btn-cancel  { background: #f5f5f5; color: #333; }
  .btn-modal:hover { opacity: .88; }
  .dur-row { display: flex; align-items: center; gap: 6px; }
  .dur-row input { text-align: center; }
  .dur-btn { padding: 6px 10px; border: 1.5px solid #d0d0d0; border-radius: 6px; background: #f5f5f5; cursor: pointer; font-size: 12px; }
  .dur-btn:hover { background: #e0e0e0; }
  .session-mgr { margin-top: 6px; }
  .session-tag {
    display: inline-flex; align-items: center; gap: 4px;
    padding: 3px 8px; border-radius: 20px; font-size: 11px;
    background: #f0f0f0; margin: 2px; cursor: default;
  }
  .session-tag .del-tag { cursor: pointer; color: #999; font-size: 13px; }
  .session-tag .del-tag:hover { color: #e53935; }
  .meta-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
  .helper-text { font-size: 11px; color: #777; margin-top: 4px; line-height: 1.4; }
</style>
</head>
<body>

<!-- TOOLBAR -->
<header class="toolbar">
  <span class="logo">&#128197;</span>
  <h1>ACT Conference 2026 &mdash; Planning Interactif</h1>
  <button class="btn btn-add" onclick="openMetaModal()">Infos</button>
  <button class="btn btn-add" onclick="openAddSessionModal()">+ Session</button>
  <button class="btn btn-sync" onclick="syncFromOnline()">Synchroniser depuis la version en ligne</button>
  <button class="btn btn-save" onclick="saveSchedule()">&#128190; Sauvegarder</button>
  <button class="btn btn-docx" onclick="generateDocx()">&#128196; Running Sheet</button>
</header>

<!-- LEGEND -->
<aside class="legend" id="legend"></aside>

<!-- MAIN -->
<main class="main">
  <div class="grid-wrapper">
    <div class="time-axis" id="time-axis"></div>
    <div class="days-row" id="days-row"></div>
  </div>
</main>

<!-- TOAST -->
<div class="toast" id="toast"></div>

<!-- EDIT MODAL -->
<div class="modal-overlay hidden" id="modal">
  <div class="modal-card">
    <h2 id="modal-h2">Événement</h2>
    <input type="hidden" id="f-id">
    <input type="hidden" id="f-day">
    <input type="hidden" id="f-session">

    <div class="form-row">
      <label>Titre</label>
      <input type="text" id="f-title" placeholder="Titre de l'événement">
    </div>
    <div class="form-row-2">
      <div class="form-row" style="margin:0">
        <label>Heure</label>
        <input type="time" id="f-time" step="300">
      </div>
      <div class="form-row" style="margin:0">
        <label>Durée (min)</label>
        <div class="dur-row">
          <button class="dur-btn" onclick="adjustDur(-15)">-15</button>
          <input type="number" id="f-duration" min="5" step="5" style="width:70px">
          <button class="dur-btn" onclick="adjustDur(15)">+15</button>
        </div>
      </div>
    </div>

    <div class="form-row-2">
      <div class="form-row" style="margin:0">
        <label>Type</label>
        <select id="f-type" onchange="onTypeChange()">
          <option value="LOUANGE">Louange</option>
          <option value="PREDICATION">Prédication</option>
          <option value="REPAS">Repas</option>
          <option value="LOGISTIQUE">Logistique</option>
          <option value="REPETITION">Répétition</option>
          <option value="REUNION">Réunion / Brief</option>
          <option value="BOOST">Boost</option>
          <option value="STE_CENE">Ste Cène</option>
          <option value="PRIERE">Prière</option>
          <option value="INTERCESSION">Intercession</option>
          <option value="FUNDRAISING">Levée de fonds</option>
          <option value="TED_PITCH">Pitch / TED</option>
          <option value="ANNONCES">Annonces</option>
          <option value="TECHNIQUE">Technique</option>
        </select>
      </div>
      <div class="form-row" style="margin:0">
        <label>Session</label>
        <select id="f-session-sel"></select>
      </div>
    </div>

    <div class="form-row">
      <label>Couleur</label>
      <div class="color-row">
        <input type="color" id="f-color">
        <div class="type-color-presets" id="color-presets"></div>
      </div>
    </div>

    <div class="form-row">
      <label>Équipes responsables <span style="color:#aaa;font-weight:400">(séparées par des virgules)</span></label>
      <input type="text" id="f-teams" placeholder="Team Louange, Team Son, …">
    </div>

    <div class="form-row">
      <label>Détails / Notes Running Sheet</label>
      <textarea id="f-details" placeholder="Détails, instructions, notes pour la Running Sheet…"></textarea>
    </div>

    <div class="modal-actions">
      <button class="btn-modal btn-danger" id="btn-delete" onclick="deleteEvent()" style="display:none;margin-right:auto">
        &#128465; Supprimer
      </button>
      <button class="btn-modal btn-cancel" onclick="closeModal()">Annuler</button>
      <button class="btn-modal btn-primary" onclick="saveEvent()">&#10003; Enregistrer</button>
    </div>
  </div>
</div>

<!-- TYPE MODAL -->
<div class="modal-overlay hidden" id="modal-type">
  <div class="modal-card">
    <h2 id="type-modal-h2">Type d'événement</h2>
    <input type="hidden" id="t-original-key">
    <div class="form-row">
      <label>Nom affiché</label>
      <input type="text" id="t-label" placeholder="Ex: Louange">
    </div>
    <div class="form-row" id="type-key-row">
      <label>Identifiant technique</label>
      <input type="text" id="t-key" placeholder="Ex: LOUANGE">
      <div class="helper-text">Utilisé pour relier les événements à ce type. Il sera généré automatiquement si vous le laissez vide.</div>
    </div>
    <div class="form-row">
      <label>Couleur</label>
      <input type="color" id="t-color">
    </div>
    <div class="modal-actions">
      <button class="btn-modal btn-cancel" onclick="closeTypeModal()">Annuler</button>
      <button class="btn-modal btn-primary" onclick="saveType()">&#10003; Enregistrer</button>
    </div>
  </div>
</div>

<!-- ADD SESSION MODAL -->
<div class="modal-overlay hidden" id="modal-session">
  <div class="modal-card">
    <h2>Ajouter / Gérer les sessions</h2>
    <div class="form-row">
      <label>Jour</label>
      <select id="fs-day"></select>
    </div>
    <div class="form-row">
      <label>Sessions existantes</label>
      <div class="session-mgr" id="fs-existing"></div>
    </div>
    <div class="form-row">
      <label>Nouveau nom de session</label>
      <input type="text" id="fs-name" placeholder="Ex: Soirée, Matinée, Après-midi…">
    </div>
    <div class="modal-actions">
      <button class="btn-modal btn-cancel" onclick="closeSessionModal()">Fermer</button>
      <button class="btn-modal btn-primary" onclick="addSession()">+ Ajouter</button>
    </div>
  </div>
</div>

<!-- META MODAL -->
<div class="modal-overlay hidden" id="modal-meta">
  <div class="modal-card">
    <h2>Informations générales de la Running Sheet</h2>
    <div class="meta-grid">
      <div class="form-row" style="margin:0">
        <label>Nom de l'événement</label>
        <input type="text" id="m-event" placeholder="ACT Conference 2026">
      </div>
      <div class="form-row" style="margin:0">
        <label>Lieu</label>
        <input type="text" id="m-lieu" placeholder="Espace 140, Rillieux la Pape">
      </div>
    </div>
    <div class="meta-grid">
      <div class="form-row" style="margin:0">
        <label>Dates installation</label>
        <input type="text" id="m-installation" placeholder="Jeudi 14 mai 2026">
      </div>
      <div class="form-row" style="margin:0">
        <label>Dates conférence</label>
        <input type="text" id="m-conference" placeholder="Vendredi 15 – Dimanche 17 mai 2026">
      </div>
    </div>
    <div class="form-row">
      <label>Dress code</label>
      <input type="text" id="m-dresscode" placeholder="Bracelet bleu pour la Dream Team">
    </div>
    <div class="form-row">
      <label>Nombre de personnes attendues</label>
      <div class="meta-grid">
        <div class="form-row" style="margin:0">
          <label>Vendredi</label>
          <input type="text" id="m-att-vendredi" placeholder="Ex: 400">
        </div>
        <div class="form-row" style="margin:0">
          <label>Samedi</label>
          <input type="text" id="m-att-samedi" placeholder="Ex: 550">
        </div>
      </div>
      <div class="form-row" style="margin-top:12px">
        <label>Dimanche</label>
        <input type="text" id="m-att-dimanche" placeholder="Ex: 450">
      </div>
      <div class="helper-text">Ces informations sont reprises en tête de la Running Sheet exportée.</div>
    </div>
    <div class="modal-actions">
      <button class="btn-modal btn-cancel" onclick="closeMetaModal()">Annuler</button>
      <button class="btn-modal btn-primary" onclick="saveMeta()">&#10003; Enregistrer</button>
    </div>
  </div>
</div>

<script src="/app.js?v=6"></script>
</body>
</html>
"""

# ─────────────────────────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return Response(HTML_TEMPLATE, mimetype="text/html; charset=utf-8")

@app.route("/app.js")
def app_js():
    with open(JS_FILE, encoding="utf-8") as f:
        return Response(f.read(), mimetype="application/javascript; charset=utf-8")

@app.route("/api/schedule", methods=["GET"])
def get_schedule():
    # Endpoint de synchronisation: Render et localhost exposent le même JSON.
    # La version locale l'utilise pour récupérer la source de vérité hébergée.
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, encoding="utf-8") as f:
            return jsonify(normalize_schedule(json.load(f)))
    return jsonify(default_copy())

@app.route("/api/schedule", methods=["POST"])
def save_schedule():
    data = normalize_schedule(request.get_json())
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return jsonify({"status": "ok"})

@app.route("/api/generate-docx", methods=["POST"])
def generate_docx_route():
    data = normalize_schedule(request.get_json())
    buf = io.BytesIO()
    build_running_sheet(data, buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="Running_Sheet_ACT_2026.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ─────────────────────────────────────────────────────────────────
# RUNNING SHEET GENERATOR
# ─────────────────────────────────────────────────────────────────
def build_running_sheet(data, output):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GREEN  = RGBColor(0x0F, 0x4C, 0x3A)
    PINK   = RGBColor(0xCF, 0x5A, 0xFD)
    BLACK  = RGBColor(0x00, 0x00, 0x00)
    FONT   = "Funnel Display"
    FALLBACK = "Calibri"

    doc = Document()
    # Page margins
    for sect in doc.sections:
        sect.top_margin    = Cm(1.5)
        sect.bottom_margin = Cm(1.5)
        sect.left_margin   = Cm(2)
        sect.right_margin  = Cm(2)

    def run(para, text, bold=False, color=BLACK, size=10, font=FONT):
        r = para.add_run(text)
        r.bold = bold
        r.font.name = font
        r.font.size = Pt(size)
        r.font.color.rgb = color
        # Also set via XML for better compatibility
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:cs'), FALLBACK)
        rPr.insert(0, rFonts)
        return r

    def para(text="", bold=False, color=BLACK, size=10, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run(p, text, bold=bold, color=color, size=size)
        return p

    def set_cell_bg(cell, hex_color):
        """Set table cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

    def set_cell_border(cell, positions=('top','bottom','left','right'), color='CCCCCC', size=4):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for pos in positions:
            border = OxmlElement(f'w:{pos}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def add_detail_paragraph(cell, text, color=BLACK, bold=False):
        p_det = cell.add_paragraph()
        p_det.paragraph_format.space_before = Pt(0)
        p_det.paragraph_format.space_after = Pt(1)
        run(p_det, text, size=9, color=color, bold=bold)

    def add_event_row(table, ev):
        row = table.add_row()
        # Time cell
        tc = row.cells[0]
        tc.width = Cm(1.8)
        set_cell_bg(tc, 'F5F5F5')
        set_cell_border(tc)
        p_time = tc.paragraphs[0]
        p_time.paragraph_format.space_before = Pt(2)
        p_time.paragraph_format.space_after  = Pt(2)
        start_time = ev.get('time', '')
        run(p_time, start_time, bold=True, size=9, color=BLACK)

        if start_time and ev.get('duration'):
            hour, minute = [int(part) for part in start_time.split(":")]
            end_total = hour * 60 + minute + int(ev.get('duration', 0))
            end_time = f"{end_total // 60:02d}:{end_total % 60:02d}"
            p_end = tc.add_paragraph()
            p_end.paragraph_format.space_before = Pt(0)
            p_end.paragraph_format.space_after = Pt(2)
            run(p_end, end_time, size=8, color=RGBColor(0x88, 0x88, 0x88))

        # Content cell
        ac = row.cells[1]
        set_cell_border(ac)
        p_act = ac.paragraphs[0]
        p_act.paragraph_format.space_before = Pt(2)
        p_act.paragraph_format.space_after  = Pt(2)

        # Title (bold, black)
        title = ev.get('title','')
        dur   = ev.get('duration', 0)
        run(p_act, title, bold=True, size=9.5, color=BLACK)
        if dur:
            run(p_act, f'  ({dur} min)', bold=False, size=8.5, color=RGBColor(0x88,0x88,0x88))

        # Details (normal, black)
        details = ev.get('details','').strip()
        if details:
            for raw_line in details.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                add_detail_paragraph(ac, line, color=BLACK)

        # Teams (pink)
        teams = [t for t in (ev.get('teams') or []) if t]
        if teams:
            add_detail_paragraph(ac, '\u25b6 ' + ' – '.join(teams), color=PINK, bold=True)

    # ── HEADER INFO ──
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    run(title, data.get('event', 'ACT Conference 2026'), bold=True, size=16, color=GREEN)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_after = Pt(2)
    run(p_h, 'Dates : ', bold=True, size=10)
    run(
        p_h,
        f"Installation : {data.get('installation_dates', '')}  |  ACT Conférence : {data.get('conference_dates', '')}",
        size=10,
    )

    p_l = doc.add_paragraph()
    p_l.paragraph_format.space_after = Pt(2)
    run(p_l, 'Lieu : ', bold=True, size=10)
    run(p_l, data.get('lieu','Espace 140, Rillieux la Pape'), size=10)

    p_dc = doc.add_paragraph()
    p_dc.paragraph_format.space_after = Pt(2)
    run(p_dc, 'Dress code : ', bold=True, size=10)
    run(p_dc, data.get('dresscode','Bracelet bleu pour la Dream Team'), size=10)

    p_att = doc.add_paragraph()
    p_att.paragraph_format.space_after = Pt(2)
    run(p_att, 'Nombre de personnes attendues : ', bold=True, size=10)

    attendance = data.get("attendance", {})
    for key, label in [("vendredi", "Vendredi 15 mai"), ("samedi", "Samedi 16 mai"), ("dimanche", "Dimanche 17 mai")]:
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_after = Pt(1)
        run(p_line, f"{label} : ", bold=True, size=10)
        run(p_line, attendance.get(key, ""), size=10)

    # ── PER DAY ──
    for day in data.get('days', []):
        doc.add_paragraph()  # spacer

        # Day header
        p_day = doc.add_paragraph()
        p_day.paragraph_format.space_before = Pt(6)
        p_day.paragraph_format.space_after  = Pt(4)
        run(p_day, day.get('name',''), bold=True, size=14, color=GREEN)

        for session in day.get('sessions', []):
            events = [e for e in session.get('events', []) if e.get('title')]
            if not events:
                continue
            events.sort(key=lambda ev: ev.get('time', '00:00'))

            # Session header
            p_sess = doc.add_paragraph()
            p_sess.paragraph_format.space_before = Pt(4)
            p_sess.paragraph_format.space_after  = Pt(2)
            run(p_sess, session.get('name',''), bold=True, size=11, color=PINK)

            # Events table
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'
            # Column widths
            for i, width in enumerate([Cm(1.8), Cm(14.5)]):
                for cell in tbl.columns[i].cells:
                    cell.width = width

            for ev in events:
                add_event_row(tbl, ev)

            doc.add_paragraph()  # spacing after table

    doc.save(output)

# ─────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5001"))
    local_url = f"http://localhost:{port}"
    print(f"\n  ACT Conference 2026 – Planning Interactif")
    print(f"  ==========================================")
    print(f"  Démarrage du serveur...")
    print(f"  En local, ouvrez votre navigateur sur: {local_url}\n")
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)